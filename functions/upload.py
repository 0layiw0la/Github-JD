import PyPDF2
from io import BytesIO
import json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH



UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'pdf'}


def build_profile_docx(user, bullets):
    """
    Generate a .docx matching the uploaded PDF's Arial-based styles,
    with centered name/contacts, full-width horizontal lines after each heading,
    proper bullet lists, and hyperlinked contacts.
    """
    doc = Document()
    styles = doc.styles
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # Title style (Arial Bold, 28pt)
    title_style = styles.add_style('ProfileTitle', 1)
    title_font = title_style.font
    title_font.name = 'Arial'
    title_font.size = Pt(28)
    title_font.bold = True
    

    # Contact style (Arial, 9pt)
    contact_style = styles.add_style('ProfileContact', 1)
    contact_font = contact_style.font
    contact_font.name = 'Arial'
    contact_font.size = Pt(9)
    contact_font.color.rgb = RGBColor(0, 0, 255)

    # Section header style (Arial Bold, 11pt)
    header_style = styles.add_style('ProfileHeader', 1)
    header_font = header_style.font
    header_font.name = 'Arial'
    header_font.size = Pt(13)
    header_font.bold = True

    # Body text style (Arial, 9pt)
    body_style = styles.add_style('ProfileBody', 1)
    body_font = body_style.font
    body_font.name = 'Arial'
    body_font.size = Pt(9)


    def insert_horizontal_line(doc):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_ALIGN_PARAGRAPH.LEFT
        table.allow_autofit = True
        cell = table.cell(0, 0)
        cell.text = ""

        # Add bottom border to the cell
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '8')       
        bottom.set(qn('w:space'), '0')
        bottom.set(qn('w:color'), '000000')
        borders.append(bottom)
        tc_pr.append(borders)

        # Remove spacing in the cell's paragraph
        para = cell.paragraphs[0]
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1

        # Reduce row height
        table.rows[0].height = Pt(0)

        # Remove cell margins by modifying table XML directly
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblCellMar = OxmlElement('w:tblCellMar')
        for m in ['top', 'left', 'bottom', 'right']:
            elem = OxmlElement(f'w:{m}')
            elem.set(qn('w:w'), '0')
            elem.set(qn('w:type'), 'dxa')
            tblCellMar.append(elem)
        tblPr.append(tblCellMar)

        doc.add_paragraph()  # This creates an empty paragraph
        para = doc.paragraphs[-1]
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(1)  # Adds 2pt space after this paragraph



    def add_hyperlink(paragraph, url, text):
        """
        Add a hyperlink to a paragraph.
        """
        part = paragraph.part
        r_id = part.relate_to(
            url,
            'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
            is_external=True
        )
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0000FF')
        rPr.append(color)
        new_run.append(rPr)
        text_elem = OxmlElement('w:t')
        text_elem.text = text
        new_run.append(text_elem)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

    # Title centered
    title_p = doc.add_paragraph(user.fullname or "Full Name Not Set", style='ProfileTitle')
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(4)

    # Contact line centered with hyperlinks
    contact_p = doc.add_paragraph(style='ProfileContact')
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_before = Pt(0)
    contact_p.paragraph_format.space_after = Pt(0)
    

    parts = []
    if user.email:
        add_hyperlink(contact_p, f"mailto:{user.email}", user.email)
        parts.append('email')
    if user.linkedin:
        if parts: contact_p.add_run(" | ")
        add_hyperlink(contact_p, user.linkedin, "LinkedIn")
        parts.append('linkedin')
    if user.github:
        if parts: contact_p.add_run(" | ")
        add_hyperlink(contact_p, user.github, "GitHub")

    # Skills
    skills_p = doc.add_paragraph("SKILLS", style='ProfileHeader')
    skills_p.paragraph_format.space_after = Pt(0)  
    insert_horizontal_line(doc)

    doc.add_paragraph(f"Programming languages: {user.programming_languages or 'N/A'}", style='ProfileBody').runs[0].bold = True
    doc.add_paragraph(f"Frameworks & Tools: {user.libraries or 'N/A'}", style='ProfileBody').runs[0].bold = True
    doc.add_paragraph(f"Machine Learning: {user.tools or 'N/A'}", style='ProfileBody').runs[0].bold = True

    # Projects
    proj_p = doc.add_paragraph("PROJECTS", style='ProfileHeader')
    proj_p.paragraph_format.space_after = Pt(0)
    insert_horizontal_line(doc)

    if bullets:
        for project, points in bullets.items():
            header = doc.add_paragraph(project.replace('-', ' ').upper(), style='ProfileBody')
            header.runs[0].bold = True
            if points:
                for pt in points:
                    p = doc.add_paragraph(pt, style='List Bullet')
                    p.style.paragraph_format.left_indent = Inches(0.25)
            else:
                p = doc.add_paragraph("No bullet points available.", style='ProfileBody')
                p.italic = True
    else:
        p = doc.add_paragraph("No projects to display.", style='ProfileBody')
        p.italic = True

    # Extracurriculars
    extra_p = doc.add_paragraph("EXTRACURRICULARS", style='ProfileHeader')
    extra_p.paragraph_format.space_after = Pt(0)
    insert_horizontal_line(doc)

    if user.extracurriculars:
        ec = json.loads(user.extracurriculars)
        for act, desc in zip(ec.get('activities', []), ec.get('descriptions', [])):
            header = doc.add_paragraph(act, style='ProfileBody')
            header.runs[0].bold = True
            if desc:
                p = doc.add_paragraph(desc, style='List Bullet')
                p.style.paragraph_format.left_indent = Inches(0.25)
    else:
        p = doc.add_paragraph("No extracurricular activities listed.", style='ProfileBody')
        p.italic = True

    # Save to BytesIO
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio




def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(pdf_path):
    text = ""
    with open(pdf_path, "rb") as pdf_file:
        reader = PyPDF2.PdfReader(pdf_file)
        for page in reader.pages:
            text += page.extract_text() + "\n"
    return text

def process_text(text):
    return text.strip()

def process_pdf(file_path):
    return extract_text_from_pdf(file_path)


