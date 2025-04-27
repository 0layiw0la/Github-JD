from io import BytesIO
import json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL


def build_profile_docx(user, bullets):
    """
    Generate a .docx matching the uploaded PDF's Arial-based styles,
    with centered name/contacts, full-width horizontal lines after each heading,
    proper bullet lists, and hyperlinked contacts.
    """
    doc = Document()
    styles = doc.styles
    section = doc.sections[0]
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    style = styles['Normal']
    style.paragraph_format.line_spacing = 1.0

    # Title style (Arial Bold, 28pt)
    title_style = styles.add_style('ProfileTitle', 1)
    title_style.paragraph_format.line_spacing = 1.0
    title_font = title_style.font
    title_font.name = 'Arial'
    title_font.size = Pt(28)
    title_font.bold = True
    

    # Contact style (Arial, 9pt)
    contact_style = styles.add_style('ProfileContact', 1)
    contact_style.paragraph_format.line_spacing = 1.0
    contact_font = contact_style.font
    contact_font.name = 'Arial'
    contact_font.size = Pt(9)
    contact_font.color.rgb = RGBColor(0, 0, 255)

    # Section header style (Arial Bold, 11pt)
    header_style = styles.add_style('ProfileHeader', 1)
    header_style.paragraph_format.line_spacing = 1.0
    header_font = header_style.font
    header_font.name = 'Arial'
    header_font.size = Pt(13)
    header_font.bold = True

    # Body text style (Arial, 9pt)
    body_style = styles.add_style('ProfileBody', 1)
    body_style.paragraph_format.line_spacing = 1.0
    body_font = body_style.font
    body_font.name = 'Arial'
    body_font.size = Pt(9)

    def insert_horizontal_line(doc):
        # create a one-cell table
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_ALIGN_PARAGRAPH.LEFT
        table.allow_autofit = True
        cell = table.cell(0, 0)

        # add bottom border
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '8')
        bottom.set(qn('w:space'), '0')
        bottom.set(qn('w:color'), '000000')
        borders.append(bottom)
        tc_pr.append(borders)

        # ensure there's a run with a non-breaking space
        para = cell.paragraphs[0]
        run = para.add_run('\u00A0')          # NBSP “text”
        run.font.size = Pt(3)                # force 3 pt
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # collapse the row height
        table.rows[0].height = Pt(2)         # try 2 pt if 0 doesn’t work

        # zero out all cell margins
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblCellMar = OxmlElement('w:tblCellMar')
        for side in ('top', 'left', 'bottom', 'right'):
            mar = OxmlElement(f'w:{side}')
            mar.set(qn('w:w'), '0')
            mar.set(qn('w:type'), 'dxa')
            tblCellMar.append(mar)
        tblPr.append(tblCellMar)

        # follow with a tiny spacer paragraph
        para2 = doc.add_paragraph('\u00A0')
        r2 = para2.runs[0]
        r2.font.size = Pt(3)
        para2.paragraph_format.space_before = Pt(0)
        para2.paragraph_format.space_after = Pt(1)

    def add_hyperlink(paragraph, url, text):
        """
        Add a hyperlink to a paragraph.
        """
        part = paragraph.part
        r_id = part.relate_to(
            url,
            'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
            is_external=True
        )
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0000FF')
        rPr.append(color)
        new_run.append(rPr)
        text_elem = OxmlElement('w:t')
        text_elem.text = text
        new_run.append(text_elem)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

    # Title centered
    title_p = doc.add_paragraph(user.fullname or "Full Name Not Set", style='ProfileTitle')
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.line_spacing = 1.0 
    title_p.paragraph_format.space_after = Pt(4)

    # Contact line centered with hyperlinks
    contact_p = doc.add_paragraph(style='ProfileContact')
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.line_spacing = 1.0
    contact_p.paragraph_format.space_before = Pt(0)
    contact_p.paragraph_format.space_after = Pt(4)
    

    parts = []
    if user.email:
        add_hyperlink(contact_p, f"mailto:{user.email}", user.email)
        parts.append('email')
    if user.linkedin:
        if parts: contact_p.add_run(" | ")
        add_hyperlink(contact_p, user.linkedin, "LinkedIn")
        parts.append('linkedin')
    if user.github:
        if parts: contact_p.add_run(" | ")
        add_hyperlink(contact_p, user.github, "GitHub")

    # Skills
    skills_p = doc.add_paragraph("SKILLS", style='ProfileHeader')
    skills_p.paragraph_format.space_after = Pt(0)  
    insert_horizontal_line(doc)

    skills_line = doc.add_paragraph(f"Programming languages: {user.programming_languages or 'N/A'}", style='ProfileBody')
    skills_line.runs[0].bold = True
    skills_line.paragraph_format.space_after = Pt(4)

    frameworks_line = doc.add_paragraph(f"Frameworks & Tools: {user.libraries or 'N/A'}", style='ProfileBody')
    frameworks_line.runs[0].bold = True
    frameworks_line.paragraph_format.space_after = Pt(4)

    tools_line = doc.add_paragraph(f"Machine Learning: {user.tools or 'N/A'}", style='ProfileBody')
    tools_line.runs[0].bold = True

    #Education 
    edu_p = doc.add_paragraph("EDUCATION", style='ProfileHeader')
    edu_p.paragraph_format.space_after = Pt(0)  
    insert_horizontal_line(doc)

    paragraph = doc.add_paragraph()
    run1 = paragraph.add_run(f"{user.school_name}")
    paragraph.add_run('                                   ' + (' '*97))
    run2 = paragraph.add_run(f"{user.duration or 'N/A'}")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph(f"{user.course or 'N/A'}", style='ProfileBody').bold = True


    # Work Experience
    work_p = doc.add_paragraph("WORK EXPERIENCE", style='ProfileHeader')
    work_p.paragraph_format.space_after = Pt(0)
    insert_horizontal_line(doc)

    # Sample work entry with justified company name and duration
    work_entry = doc.add_paragraph()
    work_company = work_entry.add_run("[Company Name]")  # Placeholder
    work_entry.add_run('                                   ' + (' '*97))  # Same spacing as education
    work_date = work_entry.add_run("[Duration]")  # Placeholder
    work_entry.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Role title
    work_role = doc.add_paragraph("[Role Title]", style='ProfileBody')
    work_role.bold = True
    work_role.paragraph_format.space_after = Pt(4)

    # Bullet point placeholder
    bullet_point = doc.add_paragraph("[Your bullet point here]", style='List Bullet')
    bullet_point.style.paragraph_format.left_indent = Inches(0.25)

    # Projects
    proj_p = doc.add_paragraph("PROJECTS", style='ProfileHeader')
    proj_p.paragraph_format.space_after = Pt(0)
    insert_horizontal_line(doc)

    if bullets:
        for project, points in bullets.items():
            header = doc.add_paragraph(style='ProfileBody')
            project_run = header.add_run(project.replace('-', ' ').upper())
            project_run.bold = True
            
            # Add spacing
            header.add_run('       ')
            
            # Add view more link
            github_url = f"https://github.com/{user.username}/{project}"
            add_hyperlink(header, github_url, "View_more")
            header.paragraph_format.space_after = Pt(4) 
            if points:
                for pt in points:
                    p = doc.add_paragraph(pt, style='List Bullet')
                    p.style.paragraph_format.left_indent = Inches(0.25)
            else:
                p = doc.add_paragraph("No bullet points available.", style='ProfileBody')
                p.italic = True
    else:
        p = doc.add_paragraph("No projects to display.", style='ProfileBody')
        p.italic = True

    # Extracurriculars
    extra_p = doc.add_paragraph("EXTRACURRICULARS", style='ProfileHeader')
    extra_p.paragraph_format.space_after = Pt(0)
    insert_horizontal_line(doc)

    if user.extracurriculars:
        ec = json.loads(user.extracurriculars)
        for act, desc in zip(ec.get('activities', []), ec.get('descriptions', [])):
            header = doc.add_paragraph(act, style='ProfileBody')
            header.runs[0].bold = True
            header.paragraph_format.space_after = Pt(4)
            if desc:
                p = doc.add_paragraph(desc, style='List Bullet')
                p.style.paragraph_format.left_indent = Inches(0.25)
    else:
        p = doc.add_paragraph("No extracurricular activities listed.", style='ProfileBody')
        p.italic = True

    # Save to BytesIO
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

